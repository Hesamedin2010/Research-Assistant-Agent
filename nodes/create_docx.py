from docx import Document
import os
from datetime import datetime
from semantic_memory import save_semantic_memory
from memory import save_memory

def publish_docx(state) -> dict:
    print("--WORD FILE GENERATION--")
    output_dir = "output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"research_{timestamp}.docx"
    pdf_path = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading('Research Report', level=1)
    doc.add_paragraph(state["edited_text"])
    doc.add_heading('References', level=1)
    for reference in state["references"]:
        doc.add_paragraph(reference)

    doc.save(pdf_path)
    print(f"DOCX saved to: {pdf_path}")

    # ✅ Store it in state before accessing
    state["pdf_path"] = pdf_path

    save_memory({
        "query": state["query"],
        "edited_text": state["edited_text"],
        "pdf_path": pdf_path,
        "references": state.get("references", [])
    })

    save_semantic_memory({
        "query": state["query"],
        "edited_text": state["edited_text"],
        "pdf_path": pdf_path,
        "references": state.get("references", [])
    })

    return state


